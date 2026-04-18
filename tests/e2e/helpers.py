"""
Shared helpers and constants for e2e tests.
"""

from pathlib import Path

from docx import Document

from evernote_to_gdrive.models import (
    AttachmentPolicy,
    MigrationOptions,
    MigrationRecord,
    MigrationStatus,
    OutputMode,
)

FIXTURES_DIR = Path(__file__).parent.parent / "input" / "sanity"
FIXTURES_EXTENDED_DIR = Path(__file__).parent.parent / "input" / "extended"


def make_local_options(dest: Path, **overrides) -> MigrationOptions:
    """Build a default local MigrationOptions, with optional field overrides."""
    defaults = dict(
        output_mode=OutputMode.LOCAL,
        dest=str(dest),
        notebooks=[],
        stacks=[],
        note=None,
        attachments=AttachmentPolicy.DOC,
        log_file=None,
        verbose=True,
    )
    defaults.update(overrides)
    return MigrationOptions(**defaults)


def assert_status_all_success(records: list[MigrationRecord]) -> None:
    failed = [r for r in records if r.status != MigrationStatus.SUCCESS]
    assert not failed, (
        f"{len(failed)} note(s) did not succeed:\n"
        + "\n".join(f"  {r.title}: {r.status} {r.error}" for r in failed)
    )


def docx_text(path: Path) -> str:
    """Return the full text of a docx file as a newline-joined string."""
    doc = Document(str(path))
    return "\n".join(p.text for p in doc.paragraphs)


def docx_external_hyperlinks(path: Path) -> list[str]:
    """Return all external hyperlink targets from a .docx file."""
    doc = Document(str(path))
    HL = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"
    return [rel.target_ref for rel in doc.part.rels.values()
            if rel.reltype == HL and rel.is_external]
