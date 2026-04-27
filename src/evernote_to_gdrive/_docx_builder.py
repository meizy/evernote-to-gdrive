"""
Build a python-docx Document from an Evernote note's ENML.

All functions here operate purely in memory — no filesystem I/O.
Sibling file writing and timestamp setting live in local_writer.py.
"""

from __future__ import annotations

import base64
import hashlib
import itertools
import logging
import re
from copy import deepcopy

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from lxml import html as lxml_html

from html4docx import HtmlToDocx
import html4docx.utils as _h4d_utils

from .classifier import _EMBEDDABLE_IMAGE_MIME, IMAGE_MAX_WIDTH_PX, _is_rtl, format_tags
from ._enml import sanitize_enml, parse_media_tag, source_url_html
from ._image import apply_exif_orientation
from .parser import Attachment, Note


# ── html4docx compatibility patches ──────────────────────────────────────────

class _SuppressHtml4docx(logging.Filter):
    def filter(self, record):
        return 'html4docx' not in (record.pathname or '')

logging.getLogger().addFilter(_SuppressHtml4docx())

# html4docx uses bare print() for unsupported-unit warnings; silence them.
def _silent_unit_converter(unit_value: str, target_unit: str = "pt"):
    unit_value = unit_value.strip().lower()
    value = float(re.sub(r"[^0-9.]", "", unit_value) or 0)
    unit = re.sub(r"[0-9.]", "", unit_value)
    from docx.shared import Pt, Cm, Inches, Mm
    from html4docx import constants
    conversion_to_pt = {
        "px": value * 0.75,
        "pt": value * 1.0,
        "in": value * 72.0,
        "pc": value * 12.0,
        "cm": value * 28.3465,
        "mm": value * 2.83465,
        "em": value * 12.0,
        "rem": value * 12.0,
        "%": value,
    }
    if unit not in conversion_to_pt:
        return None
    value_in_pt = min(conversion_to_pt[unit], constants.MAX_INDENT * 72.0)
    return {"pt": Pt(value_in_pt), "px": round(value_in_pt / 0.75, 2),
            "in": Inches(value_in_pt / 72.0), "cm": Cm(value_in_pt / 28.3465),
            "mm": Mm(value_in_pt / 2.83465)}.get(target_unit)

_h4d_utils.unit_converter = _silent_unit_converter

# html4docx bug: add_styles_to_table_cell passes parse_color()'s raw list
# directly to run.font.color.rgb, which requires an RGBColor object.
import html4docx.h4d as _h4d
_orig_add_styles_to_table_cell = _h4d.HtmlToDocx.add_styles_to_table_cell
def _patched_add_styles_to_table_cell(self, styles, doc_cell, cell_row):
    if 'color' in styles:
        from html4docx import utils as _u
        color = _u.parse_color(styles['color'])
        if color and isinstance(color, list):
            styles = dict(styles)
            styles['color'] = _u.rgb_to_hex(color)
    _orig_add_styles_to_table_cell(self, styles, doc_cell, cell_row)
_h4d.HtmlToDocx.add_styles_to_table_cell = _patched_add_styles_to_table_cell

# html4docx only treats http/https hrefs as external hyperlinks; anything else
# (including relative file paths used for local-mode inter-note links) becomes a
# broken internal w:anchor bookmark. Wrap handle_link to force any non-anchor,
# non-http href through the external-hyperlink path by using a unique sentinel URL,
# then immediately repoint the just-created relationship to the real href.
_sentinel_counter = itertools.count()
_orig_handle_link = _h4d.HtmlToDocx.handle_link


def _patched_handle_link(self, href, text, tooltip=None):
    if href and not href.startswith('#') and not href.startswith(('http:', 'https:', 'mailto:')):
        sentinel = f"http://e2g.local/{next(_sentinel_counter)}"
        _orig_handle_link(self, sentinel, text, tooltip)
        for rel in self.paragraph.part.rels.values():
            if getattr(rel, 'is_external', False) and rel._target == sentinel:
                rel._target = href
                return
        return
    _orig_handle_link(self, href, text, tooltip)


_h4d.HtmlToDocx.handle_link = _patched_handle_link


# ── docx XML helpers ──────────────────────────────────────────────────────────

def _set_para_rtl(paragraph) -> None:
    """Mark a python-docx paragraph as right-to-left.

    Also removes any paragraph-level jc=right that html4docx may have set from
    a 'text-align: right' style on the HTML element. In bidi paragraphs, renderers
    (e.g. Google Docs) treat jc=right as end-of-line (= left side for RTL), so the
    text ends up left-aligned. Without any explicit jc, bidi paragraphs default to
    start-of-line alignment (= right side for RTL), which is the desired behaviour.
    """
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)
    jc = pPr.find(qn("w:jc"))
    if jc is not None and jc.get(qn("w:val")) == "right":
        pPr.remove(jc)


def add_file_hyperlink(doc: Document, display_text: str, filename: str) -> None:
    """Add a paragraph with a clickable hyperlink to a sibling file."""
    paragraph = doc.add_paragraph()
    _append_hyperlink_to_paragraph(paragraph, display_text, filename)


def _build_hyperlink_element(paragraph, display_text: str, target: str):
    """Build a hyperlink XML element targeting a sibling file."""
    doc = paragraph.part.document

    r_id = doc.part.relate_to(
        target,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    run.append(rPr)

    text_elem = OxmlElement("w:t")
    text_elem.text = display_text
    run.append(text_elem)
    hyperlink.append(run)
    return hyperlink


def _append_hyperlink_to_paragraph(paragraph, display_text: str, target: str) -> None:
    """Append a clickable hyperlink run to an existing paragraph."""
    paragraph._p.append(_build_hyperlink_element(paragraph, display_text, target))

    if _is_rtl(display_text):
        _set_para_rtl(paragraph)


def _clear_paragraph_content(paragraph) -> None:
    """Remove paragraph content while preserving paragraph properties/style."""
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def _insert_paragraph_after(paragraph) -> Paragraph:
    """Insert a new paragraph after paragraph, preserving paragraph properties."""
    new_p = OxmlElement("w:p")
    p_pr = paragraph._p.find(qn("w:pPr"))
    if p_pr is not None:
        new_p.append(deepcopy(p_pr))
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def _split_attachment_paragraph(paragraph, token: str, display_text: str, filename: str) -> int:
    """Split a paragraph around a standalone token run and insert a hyperlink paragraph."""
    content = [child for child in paragraph._p if child.tag != qn("w:pPr")]
    run_text_by_element = {id(run._r): run.text for run in paragraph.runs}
    token_index = None
    for idx, child in enumerate(content):
        if child.tag == qn("w:r") and run_text_by_element.get(id(child)) == token:
            token_index = idx
            break
    if token_index is None:
        return 0

    before = content[:token_index]
    token_node = content[token_index]
    after = content[token_index + 1:]

    if not before:
        paragraph._p.remove(token_node)
        _clear_paragraph_content(paragraph)
        _append_hyperlink_to_paragraph(paragraph, display_text, filename)
        hyperlink_paragraph = paragraph
    else:
        paragraph._p.remove(token_node)
        hyperlink_paragraph = _insert_paragraph_after(paragraph)
        _append_hyperlink_to_paragraph(hyperlink_paragraph, display_text, filename)

    for child in after:
        paragraph._p.remove(child)
    if after:
        after_paragraph = _insert_paragraph_after(hyperlink_paragraph)
        for child in after:
            after_paragraph._p.append(child)
    return 1


def replace_attachment_placeholder(doc: Document, token: str, display_text: str, filename: str) -> int:
    """Replace placeholder paragraphs with in-place file hyperlinks."""
    replaced = 0
    for para in doc.paragraphs:
        if para.text.strip() != token:
            replaced += _split_attachment_paragraph(para, token, display_text, filename)
            continue
        _clear_paragraph_content(para)
        _append_hyperlink_to_paragraph(para, display_text, filename)
        replaced += 1
    return replaced


# ── ENML → HTML ───────────────────────────────────────────────────────────────

def attachment_hash_map(attachments: list[Attachment]) -> dict[str, Attachment]:
    """Build MD5-hex → Attachment map (matches the hash attr in ENML <en-media>)."""
    return {att.hash: att for att in attachments}


def _enml_to_docx_html(enml: str, hash_map: dict[str, Attachment], title: str = "") -> str:
    """Convert ENML to clean HTML for html4docx.

    - <en-media> for embeddable images → <img src="data:...;base64,...">
    - <en-media> for other types (PDF etc.) → removed (handled as sibling files)
    - ENML-specific tags stripped; standard HTML preserved.
    """
    placeholder_map: dict[str, Attachment] = {}

    def _replace(m: re.Match) -> str:
        tag = m.group(0)
        h, mime = parse_media_tag(tag)
        if not h or not mime:
            return ''
        att = hash_map.get(h)
        if att and mime in _EMBEDDABLE_IMAGE_MIME:
            img_data = apply_exif_orientation(att.data, mime)
            b64 = base64.b64encode(img_data).decode()
            width_match = re.search(r'\bwidth="(\d+)"', tag)
            w = int(width_match.group(1)) if width_match else None
            if w is None or w > IMAGE_MAX_WIDTH_PX:
                w = IMAGE_MAX_WIDTH_PX
            return f'<img src="data:{mime};base64,{b64}" width="{w}px"/>'
        if att is not None:
            token = f"[[ATTACHMENT:{hashlib.md5(f'{h}:{mime}'.encode()).hexdigest()}]]"
            placeholder_map[token] = att
            return token
        return ''  # non-image or unknown hash — drop it

    html = sanitize_enml(enml, _replace, title=title)
    _enml_to_docx_html.last_placeholder_map = placeholder_map
    return html

# ── document post-processing ──────────────────────────────────────────────────

def _compact_doc_spacing(doc: Document) -> None:
    """Remove default paragraph space-after from the document defaults.

    The python-docx template ships with after=200 twips (10 pt) and line=276
    (1.15×). Evernote renders its ENML as plain HTML where <div> blocks have no
    inter-paragraph gap. Override the docDefaults so the DOCX matches.
    """
    styles_elem = doc.part.styles._element
    doc_defaults = styles_elem.find(qn("w:docDefaults"))
    if doc_defaults is None:
        return
    pPr_default = doc_defaults.find(qn("w:pPrDefault"))
    if pPr_default is None:
        return
    pPr = pPr_default.find(qn("w:pPr"))
    if pPr is None:
        return
    spacing = pPr.find(qn("w:spacing"))
    if spacing is not None:
        spacing.set(qn("w:after"), "0")
        spacing.set(qn("w:line"), "240")
        spacing.set(qn("w:lineRule"), "auto")


def build_html(note: Note, hash_map: dict[str, Attachment],
               include_tags: bool = True, include_source_url: bool = True) -> str:
    """Assemble the HTML string to feed into html4docx."""
    parts = []
    if include_tags and note.tags:
        parts.append(f'<p>{format_tags(note.tags)}</p>')
    if include_source_url and note.source_url:
        parts.append(source_url_html(note.source_url))
    if note.enml:
        parts.append(_enml_to_docx_html(note.enml, hash_map, title=note.title))
    return '\n'.join(parts)


def _postprocess_paragraphs(doc: Document) -> None:
    """Mark RTL paragraphs, center image-only paragraphs, mirror szCs for bidi runs."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    for para in doc.paragraphs:
        if _is_rtl(para.text):
            _set_para_rtl(para)
        # Center paragraphs that contain only an inline image and no text
        if not para.text.strip() and any(r._r.find('.//' + qn('a:graphicData')) is not None for r in para.runs):
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Mirror <w:sz> → <w:szCs> and <w:b> → <w:bCs> for complex/bidi scripts
        # (Hebrew, Arabic). Word and Google Docs use the *Cs variants for RTL runs;
        # html4docx only sets the Latin variants, so bold and font-size are ignored
        # for Hebrew text without this fix.
        for run in para.runs:
            rpr = run._r.find(qn('w:rPr'))
            if rpr is None:
                continue
            sz = rpr.find(qn('w:sz'))
            if sz is not None:
                szcs = rpr.find(qn('w:szCs'))
                if szcs is None:
                    szcs = OxmlElement('w:szCs')
                    szcs.set(qn('w:val'), sz.get(qn('w:val')))
                    sz.addnext(szcs)
                else:
                    szcs.set(qn('w:val'), sz.get(qn('w:val')))
            b = rpr.find(qn('w:b'))
            if b is not None and rpr.find(qn('w:bCs')) is None:
                bcs = OxmlElement('w:bCs')
                b.addnext(bcs)
            i = rpr.find(qn('w:i'))
            if i is not None and rpr.find(qn('w:iCs')) is None:
                ics = OxmlElement('w:iCs')
                i.addnext(ics)


def _paragraph_has_image(paragraph) -> bool:
    return any(r._r.find('.//' + qn('a:graphicData')) is not None for r in paragraph.runs)


def _paragraph_style_id(paragraph) -> str | None:
    """Return the paragraph style id from XML without triggering style resolution."""
    ppr = paragraph._p.pPr
    if ppr is None or ppr.pStyle is None:
        return None
    return ppr.pStyle.val


def _is_list_paragraph(paragraph) -> bool:
    style_id = _paragraph_style_id(paragraph)
    return bool(style_id and style_id.startswith(("ListBullet", "ListNumber")))


def _is_normal_paragraph(paragraph) -> bool:
    style_id = _paragraph_style_id(paragraph)
    return style_id in (None, "Normal")


def _strip_terminal_break_runs(doc: Document) -> None:
    """Remove paragraph-final line breaks that html4docx emits from trailing <br> tags."""
    for para in doc.paragraphs:
        while para.runs:
            run = para.runs[-1]
            r = run._r
            children = list(r)
            removed = False
            while children and children[-1].tag in {qn('w:br'), qn('w:cr')}:
                r.remove(children[-1])
                children.pop()
                removed = True
            if len(r) == 0:
                para._p.remove(r)
                continue
            if not removed:
                break


def _strip_leading_break_runs(doc: Document) -> None:
    """Remove paragraph-leading line breaks that render as a blank first line."""
    for para in doc.paragraphs:
        while True:
            children = list(para._p)
            content_children = [child for child in children if child.tag != qn('w:pPr')]
            if not content_children or content_children[0].tag != qn('w:r'):
                break
            first_run = content_children[0]
            run_children = list(first_run)
            removed = False
            while run_children and run_children[0].tag in {qn('w:br'), qn('w:cr')}:
                first_run.remove(run_children[0])
                run_children.pop(0)
                removed = True
            if len(first_run) == 0:
                para._p.remove(first_run)
                continue
            if not removed:
                break


def _strip_empty_list_paragraphs(doc: Document) -> None:
    """Drop empty list paragraphs; Evernote often uses them as editor spacers only."""
    for para in list(doc.paragraphs):
        if not _is_list_paragraph(para):
            continue
        if para.text.strip() or _paragraph_has_image(para):
            continue
        para._element.getparent().remove(para._element)


def _insert_break_before_continuation_urls(doc: Document) -> None:
    """Restore line breaks before bare URL continuations inside list paragraphs."""
    for para in doc.paragraphs:
        if not _is_list_paragraph(para):
            continue
        for child in list(para._p):
            if child.tag != qn("w:hyperlink"):
                continue
            texts = child.findall('.//' + qn('w:t'))
            url_text = "".join(t.text or "" for t in texts).strip()
            if not re.match(r"https?://", url_text):
                continue
            prev = child.getprevious()
            if prev is None:
                continue
            if prev.tag == qn("w:r") and prev.findall(qn("w:br")):
                continue
            br_run = OxmlElement("w:r")
            br_run.append(OxmlElement("w:br"))
            para._p.insert(para._p.index(child), br_run)


def _strip_heading_to_list_separators(doc: Document) -> None:
    """Drop empty Normal paragraphs inserted between a heading-like line and its list."""
    paras = list(doc.paragraphs)
    for idx, para in enumerate(paras[1:-1], start=1):
        if not _is_normal_paragraph(para) or para.text.strip() or _paragraph_has_image(para):
            continue
        prev_para = paras[idx - 1]
        next_para = paras[idx + 1]
        prev_is_normal = _is_normal_paragraph(prev_para)
        next_is_list = _is_list_paragraph(next_para)
        if not prev_is_normal or not prev_para.text.strip():
            continue
        if not next_is_list:
            continue
        para._element.getparent().remove(para._element)


def _cap_image_sizes(doc: Document) -> None:
    """Resize inline images that exceed page width (500 px) or height (700 px)."""
    _MAX_CX = int(500 * 9525)   # 500 px in EMU  (1 px = 914400/96 = 9525 EMU at 96 dpi)
    _MAX_CY = int(700 * 9525)   # 700 px in EMU
    for para in doc.paragraphs:
        for run in para.runs:
            for inline in run._r.findall('.//' + qn('wp:inline')):
                extent = inline.find(qn('wp:extent'))
                if extent is None:
                    continue
                cx = int(extent.get('cx', '0'))
                cy = int(extent.get('cy', '0'))
                if cx <= 0 or cy <= 0:
                    continue
                scale = min(_MAX_CX / cx, _MAX_CY / cy, 1.0)
                if scale < 1.0:
                    new_cx, new_cy = int(cx * scale), int(cy * scale)
                    extent.set('cx', str(new_cx))
                    extent.set('cy', str(new_cy))
                    for a_ext in inline.findall('.//' + qn('a:ext')):
                        a_ext.set('cx', str(new_cx))
                        a_ext.set('cy', str(new_cy))


def _strip_trailing_blanks(doc: Document) -> None:
    """Remove trailing empty paragraphs (no text, no image) from the document."""
    while doc.paragraphs:
        last = doc.paragraphs[-1]
        has_image = _paragraph_has_image(last)
        if last.text.strip() or has_image:
            break
        last._element.getparent().remove(last._element)



# ── public API ────────────────────────────────────────────────────────────────

def _html_to_doc(html: str, space_after_pt: int | None = None) -> Document:
    """Shared pipeline: HTML string → post-processed Document."""
    from docx.shared import Pt
    doc = Document()
    _compact_doc_spacing(doc)
    if html.strip():
        HtmlToDocx().add_html_to_document(html, doc)
    _strip_leading_break_runs(doc)
    _strip_terminal_break_runs(doc)
    _strip_empty_list_paragraphs(doc)
    _insert_break_before_continuation_urls(doc)
    _strip_leading_break_runs(doc)
    _strip_heading_to_list_separators(doc)
    _postprocess_paragraphs(doc)
    if space_after_pt is not None:
        for para in doc.paragraphs:
            if para.style.name == 'Normal':
                para.paragraph_format.space_after = Pt(space_after_pt)
    _cap_image_sizes(doc)
    _strip_trailing_blanks(doc)
    return doc


def build_doc(note: Note, attachments: list[Attachment], include_tags: bool = True) -> Document:
    """Build a python-docx Document from the note's ENML. No filesystem I/O."""
    hash_map = attachment_hash_map(attachments)
    html = build_html(note, hash_map, include_tags=include_tags)
    doc = _html_to_doc(html)
    doc._attachment_placeholders = dict(getattr(_enml_to_docx_html, "last_placeholder_map", {}))
    return doc


def build_html_doc(html: str, space_after_pt: int | None = None) -> Document:
    """Build a python-docx Document from pre-rendered HTML (e.g. Readability output)."""
    return _html_to_doc(html, space_after_pt=space_after_pt)
