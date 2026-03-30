"""
Build a python-docx Document from an Evernote note's ENML.

All functions here operate purely in memory — no filesystem I/O.
Sibling file writing and timestamp setting live in local_writer.py.
"""

from __future__ import annotations

import base64
import hashlib
import logging
import re
import sys

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from html4docx import HtmlToDocx
import html4docx.utils as _h4d_utils

from .classifier import _EMBEDDABLE_IMAGE_MIME, _is_rtl
from .parser import Attachment, Note


# ── html4docx compatibility patches ──────────────────────────────────────────

class _SuppressHtml4docx(logging.Filter):
    def filter(self, record):
        return 'html4docx' not in (record.pathname or '')

logging.getLogger().addFilter(_SuppressHtml4docx())

# html4docx uses bare print() for unsupported-unit warnings; silence them.
def _silent_unit_converter(unit_value: str, target_unit: str = "pt"):
    unit_value = unit_value.strip().lower()
    value = float(re.sub(r"[^0-9.]", "", unit_value) or 0)
    unit = re.sub(r"[0-9.]", "", unit_value)
    from docx.shared import Pt, Cm, Inches, Mm
    from html4docx import constants
    conversion_to_pt = {
        "px": value * 0.75,
        "pt": value * 1.0,
        "in": value * 72.0,
        "pc": value * 12.0,
        "cm": value * 28.3465,
        "mm": value * 2.83465,
        "em": value * 12.0,
        "rem": value * 12.0,
        "%": value,
    }
    if unit not in conversion_to_pt:
        return None
    value_in_pt = min(conversion_to_pt[unit], constants.MAX_INDENT * 72.0)
    return {"pt": Pt(value_in_pt), "px": round(value_in_pt / 0.75, 2),
            "in": Inches(value_in_pt / 72.0), "cm": Cm(value_in_pt / 28.3465),
            "mm": Mm(value_in_pt / 2.83465)}.get(target_unit)

_h4d_utils.unit_converter = _silent_unit_converter

# html4docx bug: add_styles_to_table_cell passes parse_color()'s raw list
# directly to run.font.color.rgb, which requires an RGBColor object.
import html4docx.h4d as _h4d
_orig_add_styles_to_table_cell = _h4d.HtmlToDocx.add_styles_to_table_cell
def _patched_add_styles_to_table_cell(self, styles, doc_cell, cell_row):
    if 'color' in styles:
        from html4docx import utils as _u
        color = _u.parse_color(styles['color'])
        if color and isinstance(color, list):
            styles = dict(styles)
            styles['color'] = _u.rgb_to_hex(color)
    _orig_add_styles_to_table_cell(self, styles, doc_cell, cell_row)
_h4d.HtmlToDocx.add_styles_to_table_cell = _patched_add_styles_to_table_cell


# ── docx XML helpers ──────────────────────────────────────────────────────────

def _set_para_rtl(paragraph) -> None:
    """Mark a python-docx paragraph as right-to-left.

    Also removes any paragraph-level jc=right that html4docx may have set from
    a 'text-align: right' style on the HTML element. In bidi paragraphs, renderers
    (e.g. Google Docs) treat jc=right as end-of-line (= left side for RTL), so the
    text ends up left-aligned. Without any explicit jc, bidi paragraphs default to
    start-of-line alignment (= right side for RTL), which is the desired behaviour.
    """
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)
    jc = pPr.find(qn("w:jc"))
    if jc is not None and jc.get(qn("w:val")) == "right":
        pPr.remove(jc)


def add_file_hyperlink(doc: Document, display_text: str, filename: str) -> None:
    """Add a paragraph with a clickable hyperlink to a sibling file."""
    paragraph = doc.add_paragraph()

    r_id = doc.part.relate_to(
        filename,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    run.append(rPr)

    text_elem = OxmlElement("w:t")
    text_elem.text = display_text
    run.append(text_elem)
    hyperlink.append(run)

    paragraph._p.append(hyperlink)

    if _is_rtl(display_text):
        _set_para_rtl(paragraph)


# ── ENML → HTML ───────────────────────────────────────────────────────────────

def _attachment_hash_map(attachments: list[Attachment]) -> dict[str, Attachment]:
    """Build MD5-hex → Attachment map (matches the hash attr in ENML <en-media>)."""
    return {hashlib.md5(att.data).hexdigest(): att for att in attachments}


def _strip_navigation_elements(html: str, _debug_title: str = "") -> str:
    """Remove site-chrome elements captured by the Evernote web clipper.

    The web clipper sometimes captures the full page including navigation
    menus, headers, and footers. Strip semantic tags (<nav>, <header>,
    <footer>) and any block with role="navigation" or role="banner".
    Uses lxml for robustness; falls back to regex on parse failure.
    """
    import os
    _debug = os.environ.get("EN_DEBUG_NAV")
    try:
        from lxml import etree
        import lxml.html as lxmlhtml
        root = lxmlhtml.fragment_fromstring(html, create_parent='div')
        _NAV_ROLES = {'navigation', 'banner', 'contentinfo'}
        removed: list[str] = []
        for el in list(root.iter()):
            tag = el.tag if isinstance(el.tag, str) else ''
            role = (el.get('role') or '').strip().lower()
            if tag.lower() in ('nav', 'header', 'footer') or role in _NAV_ROLES:
                parent = el.getparent()
                if parent is not None:
                    if _debug:
                        snippet = (el.text_content() if hasattr(el, 'text_content') else '')[:80].replace('\n', ' ')
                        removed.append(f"<{tag} role={role!r}> {snippet!r}")
                    if el.tail:
                        prev = el.getprevious()
                        if prev is not None:
                            prev.tail = (prev.tail or '') + el.tail
                        else:
                            parent.text = (parent.text or '') + el.tail
                    parent.remove(el)
        if _debug and removed:
            prefix = f"[nav-strip] {_debug_title!r}: " if _debug_title else "[nav-strip] "
            for r in removed:
                print(f"{prefix}removed {r}")
        elif _debug:
            prefix = f"[nav-strip] {_debug_title!r}: " if _debug_title else "[nav-strip] "
            print(f"{prefix}nothing removed")
        parts = []
        if root.text:
            parts.append(root.text)
        for child in root:
            parts.append(etree.tostring(child, encoding='unicode', method='html'))
        return ''.join(parts)
    except Exception:
        for tag in ('nav', 'header', 'footer'):
            html = re.sub(rf'<{tag}\b[^>]*>.*?</{tag}>', '', html, flags=re.DOTALL | re.IGNORECASE)
        return html


def _sanitize_enml(enml: str, hash_map: dict[str, Attachment], title: str = "") -> str:
    """Convert ENML to clean HTML for html4docx.

    - <en-media> for embeddable images → <img src="data:...;base64,...">
    - <en-media> for other types (PDF etc.) → removed (handled as sibling files)
    - ENML-specific tags stripped; standard HTML preserved.
    """
    html = re.sub(r'<\?xml[^>]*\?>', '', enml)
    html = re.sub(r'<!DOCTYPE[^>]*>', '', html)
    # Strip navigation elements BEFORE generating data URIs — lxml's HTML parser
    # truncates very long attribute values, which corrupts multi-MB base64 strings.
    html = _strip_navigation_elements(html, _debug_title=title)

    def replace_en_media(m: re.Match) -> str:
        tag = m.group(0)
        hash_match = re.search(r'hash="([0-9a-f]+)"', tag)
        type_match = re.search(r'type="([^"]+)"', tag)
        if not hash_match or not type_match:
            return ''
        h, mime = hash_match.group(1), type_match.group(1)
        att = hash_map.get(h)
        if att and mime in _EMBEDDABLE_IMAGE_MIME:
            b64 = base64.b64encode(att.data).decode()
            width_match = re.search(r'\bwidth="(\d+)"', tag)
            w = int(width_match.group(1)) if width_match else None
            _MAX_WIDTH = 500  # px — fits a standard docx page with margins
            if w is None or w > _MAX_WIDTH:
                w = _MAX_WIDTH
            return f'<img src="data:{mime};base64,{b64}" width="{w}px"/>'
        return ''  # non-image or unknown hash — drop it

    html = re.sub(r'<en-media\b[^>]*/>', replace_en_media, html)
    html = re.sub(r'<en-media\b[^>]*>.*?</en-media>', replace_en_media, html, flags=re.DOTALL)
    # Remove <img> tags with external URLs — python-docx cannot fetch remote images,
    # and web-clipped notes sometimes reference Evernote's CDN instead of embedding data.
    external_imgs = re.findall(r'<img\b[^>]*\bsrc="(https?://[^"]*)"[^>]*/?>', html, flags=re.IGNORECASE)
    if external_imgs:
        note_label = f" {title!r}" if title else ""
        print(f"WARNING:{note_label}: {len(external_imgs)} external image(s) skipped (not retrievable)", file=sys.stderr, flush=True)
    html = re.sub(r'<img\b[^>]*\bsrc="https?://[^"]*"[^>]*/>', '', html, flags=re.IGNORECASE)
    html = re.sub(r'<img\b[^>]*\bsrc="https?://[^"]*"[^>]*>', '', html, flags=re.IGNORECASE)
    # Remove encrypted blocks
    html = re.sub(r'<en-crypt\b[^>]*>.*?</en-crypt>', '', html, flags=re.DOTALL)
    # Convert checkboxes to text markers
    html = re.sub(r'<en-todo\b[^>]*checked="true"[^>]*/>', '[x]\u00a0', html)
    html = re.sub(r'<en-todo\b[^>]*/>', '[\u00a0]\u00a0', html)
    # Strip the ENML root tag (unknown tags confuse html4docx)
    html = re.sub(r'</?en-note\b[^>]*>', '', html)
    return html.strip()


# ── document post-processing ──────────────────────────────────────────────────

def _compact_doc_spacing(doc: Document) -> None:
    """Remove default paragraph space-after from the document defaults.

    The python-docx template ships with after=200 twips (10 pt) and line=276
    (1.15×). Evernote renders its ENML as plain HTML where <div> blocks have no
    inter-paragraph gap. Override the docDefaults so the DOCX matches.
    """
    styles_elem = doc.part.styles._element
    doc_defaults = styles_elem.find(qn("w:docDefaults"))
    if doc_defaults is None:
        return
    pPr_default = doc_defaults.find(qn("w:pPrDefault"))
    if pPr_default is None:
        return
    pPr = pPr_default.find(qn("w:pPr"))
    if pPr is None:
        return
    spacing = pPr.find(qn("w:spacing"))
    if spacing is not None:
        spacing.set(qn("w:after"), "0")
        spacing.set(qn("w:line"), "240")
        spacing.set(qn("w:lineRule"), "auto")


def _build_html(note: Note, hash_map: dict[str, Attachment]) -> str:
    """Assemble the HTML string to feed into html4docx."""
    parts = []
    if note.source_url:
        url = note.source_url
        parts.append(f'<p>Source: <a href="{url}">{url}</a></p>')
    if note.enml:
        parts.append(_sanitize_enml(note.enml, hash_map, title=note.title))
    return '\n'.join(parts)


def _postprocess_paragraphs(doc: Document) -> None:
    """Mark RTL paragraphs, center image-only paragraphs, mirror szCs for bidi runs."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    for para in doc.paragraphs:
        if _is_rtl(para.text):
            _set_para_rtl(para)
        # Center paragraphs that contain only an inline image and no text
        if not para.text.strip() and any(r._r.find('.//' + qn('a:graphicData')) is not None for r in para.runs):
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Mirror <w:sz> into <w:szCs> on every run that has an explicit font size.
        # LibreOffice (and Word) use szCs for complex/bidi scripts (Hebrew, Arabic).
        # html4docx only sets sz; without a matching szCs the run inherits the
        # document default (11 pt), so Hebrew text always appears at 11 pt regardless
        # of the sz value.
        for run in para.runs:
            rpr = run._r.find(qn('w:rPr'))
            if rpr is None:
                continue
            sz = rpr.find(qn('w:sz'))
            if sz is None:
                continue
            szcs = rpr.find(qn('w:szCs'))
            if szcs is None:
                szcs = OxmlElement('w:szCs')
                szcs.set(qn('w:val'), sz.get(qn('w:val')))
                sz.addnext(szcs)
            else:
                szcs.set(qn('w:val'), sz.get(qn('w:val')))


def _cap_image_sizes(doc: Document) -> None:
    """Resize inline images that exceed page width (500 px) or height (700 px)."""
    _MAX_CX = int(500 * 9525)   # 500 px in EMU  (1 px = 914400/96 = 9525 EMU at 96 dpi)
    _MAX_CY = int(700 * 9525)   # 700 px in EMU
    for para in doc.paragraphs:
        for run in para.runs:
            for inline in run._r.findall('.//' + qn('wp:inline')):
                extent = inline.find(qn('wp:extent'))
                if extent is None:
                    continue
                cx = int(extent.get('cx', '0'))
                cy = int(extent.get('cy', '0'))
                if cx <= 0 or cy <= 0:
                    continue
                scale = min(_MAX_CX / cx, _MAX_CY / cy, 1.0)
                if scale < 1.0:
                    new_cx, new_cy = int(cx * scale), int(cy * scale)
                    extent.set('cx', str(new_cx))
                    extent.set('cy', str(new_cy))
                    for a_ext in inline.findall('.//' + qn('a:ext')):
                        a_ext.set('cx', str(new_cx))
                        a_ext.set('cy', str(new_cy))


def _strip_trailing_blanks(doc: Document) -> None:
    """Remove trailing empty paragraphs (no text, no image) from the document."""
    while doc.paragraphs:
        last = doc.paragraphs[-1]
        has_image = any(r._r.find('.//' + qn('a:graphicData')) is not None for r in last.runs)
        if last.text.strip() or has_image:
            break
        last._element.getparent().remove(last._element)


# ── public API ────────────────────────────────────────────────────────────────

def build_doc(note: Note, attachments: list[Attachment]) -> Document:
    """Build a python-docx Document from the note's ENML. No filesystem I/O."""
    doc = Document()
    _compact_doc_spacing(doc)
    hash_map = _attachment_hash_map(attachments)
    html = _build_html(note, hash_map)
    if html.strip():
        HtmlToDocx().add_html_to_document(html, doc)
    _postprocess_paragraphs(doc)
    _cap_image_sizes(doc)
    _strip_trailing_blanks(doc)
    return doc
