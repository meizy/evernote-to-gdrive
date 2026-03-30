"""
Local output mode: write notes to a folder/subfolder tree on disk.

  attachment-only, single  → raw file (<title>.<ext>)
  attachment-only, multi   → one .docx listing all attachments (doc mode)
                             OR one raw file per attachment (files mode)
  text-only                → <title>.docx
  text + attachments       → <title>.docx  (images embedded, PDFs as sibling files)

RTL (Hebrew/Arabic) paragraphs are detected and marked as bidi in the .docx XML
so that Word, LibreOffice, and Google Docs all render them correctly.
"""

from __future__ import annotations

import base64
import ctypes
import ctypes.util
import ctypes.wintypes
import hashlib
import io
import os
import platform
import struct
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
import logging

from html4docx import HtmlToDocx
import html4docx.utils as _h4d_utils

logging.getLogger('root').setLevel(logging.ERROR)

# html4docx uses bare print() for unsupported-unit warnings; silence them.
def _silent_unit_converter(unit_value: str, target_unit: str = "pt"):
    import re
    unit_value = unit_value.strip().lower()
    value = float(re.sub(r"[^0-9.]", "", unit_value) or 0)
    unit = re.sub(r"[0-9.]", "", unit_value)
    from docx.shared import Pt, Cm, Inches, Mm
    from html4docx import constants
    conversion_to_pt = {
        "px": value * 0.75, "pt": value, "in": value * 72.0,
        "pc": value * 12.0, "cm": value * 28.3465, "mm": value * 2.83465,
        "em": value * 12.0, "rem": value * 12.0, "%": value,
    }
    if unit not in conversion_to_pt:
        return None  # silently return None instead of printing
    value_in_pt = min(conversion_to_pt[unit], constants.MAX_INDENT * 72.0)
    return {"pt": Pt(value_in_pt), "px": round(value_in_pt / 0.75, 2),
            "in": Inches(value_in_pt / 72.0), "cm": Cm(value_in_pt / 28.3465),
            "mm": Mm(value_in_pt / 2.83465)}.get(target_unit)

_h4d_utils.unit_converter = _silent_unit_converter

from .classifier import (
    NoteKind, attachment_drive_filename, ClassifiedNote,
    _EMBEDDABLE_IMAGE_MIME, _is_rtl, _safe_name, _ext_for_mime,
)
from .parser import Attachment, Note


def _set_para_rtl(paragraph) -> None:
    """Mark a python-docx paragraph as right-to-left.

    Also removes any paragraph-level jc=right that html4docx may have set from
    a 'text-align: right' style on the HTML element. In bidi paragraphs, renderers
    (e.g. Google Docs) treat jc=right as end-of-line (= left side for RTL), so the
    text ends up left-aligned. Without any explicit jc, bidi paragraphs default to
    start-of-line alignment (= right side for RTL), which is the desired behaviour.
    """
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)
    jc = pPr.find(qn("w:jc"))
    if jc is not None and jc.get(qn("w:val")) == "right":
        pPr.remove(jc)


def _add_file_hyperlink(doc: Document, display_text: str, filename: str) -> None:
    """Add a paragraph with a clickable hyperlink to a sibling file."""
    paragraph = doc.add_paragraph()

    r_id = doc.part.relate_to(
        filename,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    run.append(rPr)

    text_elem = OxmlElement("w:t")
    text_elem.text = display_text
    run.append(text_elem)
    hyperlink.append(run)

    paragraph._p.append(hyperlink)

    if _is_rtl(display_text):
        _set_para_rtl(paragraph)


# ── filesystem timestamps ─────────────────────────────────────────────────────

def _set_timestamps(path: Path, created: datetime | None, updated: datetime | None) -> None:
    """
    Set file timestamps to match the original Evernote note dates.
      mtime → note's updated date (fallback: created)
      birth time (macOS only) → note's created date
    """
    mtime_dt = updated or created
    if mtime_dt:
        mtime = mtime_dt.timestamp()
        os.utime(path, (mtime, mtime))

    system = platform.system()
    if system == "Darwin" and created:
        _set_macos_birthtime(path, created)
    elif system == "Windows" and created:
        _set_windows_birthtime(path, created)


def _set_macos_birthtime(path: Path, dt: datetime) -> None:
    """Set macOS file creation (birth) time via setattrlist syscall."""
    try:
        libc = ctypes.CDLL(ctypes.util.find_library("c"), use_errno=True)

        # struct attrlist { u_short bitmapcount; u_short reserved; attrgroup_t commonattr; ... }
        # ATTR_BIT_MAP_COUNT = 5, ATTR_CMN_CRTIME = 0x00000200
        attrlist_buf = struct.pack("HHiiii", 5, 0, 0x00000200, 0, 0, 0)

        # struct timespec { time_t tv_sec; long tv_nsec; }
        ts_buf = struct.pack("ll", int(dt.timestamp()), 0)

        libc.setattrlist(
            str(path).encode("utf-8"),
            ctypes.c_char_p(attrlist_buf),
            ctypes.c_char_p(ts_buf),
            ctypes.c_size_t(len(ts_buf)),
            ctypes.c_ulong(0),
        )
    except Exception:
        pass


def _set_windows_birthtime(path: Path, dt: datetime) -> None:
    """Set Windows file creation time via SetFileTime (kernel32)."""
    try:
        # Windows FILETIME: 100-nanosecond intervals since 1601-01-01
        EPOCH_DIFF = 116444736000000000  # offset between 1601 and 1970 in 100ns units
        filetime = int(dt.timestamp() * 10_000_000) + EPOCH_DIFF

        kernel32 = ctypes.windll.kernel32

        # Open file with GENERIC_WRITE, share all, no inherit, OPEN_EXISTING
        handle = kernel32.CreateFileW(
            str(path),
            0x40000000,   # GENERIC_WRITE
            0x00000007,   # FILE_SHARE_READ | FILE_SHARE_WRITE | FILE_SHARE_DELETE
            None,
            3,            # OPEN_EXISTING
            0x80,         # FILE_ATTRIBUTE_NORMAL
            None,
        )
        if handle == ctypes.wintypes.HANDLE(-1).value:
            return

        ft = ctypes.wintypes.FILETIME(filetime & 0xFFFFFFFF, filetime >> 32)
        kernel32.SetFileTime(handle, ctypes.byref(ft), None, None)
        kernel32.CloseHandle(handle)
    except Exception:
        pass


# ── folder layout ─────────────────────────────────────────────────────────────

def note_folder(output_dir: Path, note: Note) -> Path:
    parts = [output_dir]
    if note.stack:
        parts.append(note.stack)
    parts.append(note.notebook)
    folder = Path(*parts)
    folder.mkdir(parents=True, exist_ok=True)
    return folder


def _unique_path(path: Path) -> Path:
    """If path exists, append (2), (3), ... until unique."""
    if not path.exists():
        return path
    stem, suffix = path.stem, path.suffix
    n = 2
    while True:
        candidate = path.with_name(f"{stem} ({n}){suffix}")
        if not candidate.exists():
            return candidate
        n += 1


# ── docx builders ─────────────────────────────────────────────────────────────

def _attachment_hash_map(attachments: list[Attachment]) -> dict[str, Attachment]:
    """Build MD5-hex → Attachment map (matches the hash attr in ENML <en-media>)."""
    return {hashlib.md5(att.data).hexdigest(): att for att in attachments}


def _sanitize_enml(enml: str, hash_map: dict[str, Attachment]) -> str:
    """Convert ENML to clean HTML for html4docx.

    - <en-media> for embeddable images → <img src="data:...;base64,...">
    - <en-media> for other types (PDF etc.) → removed (handled as sibling files)
    - ENML-specific tags stripped; standard HTML preserved.
    """
    html = re.sub(r'<\?xml[^>]*\?>', '', enml)
    html = re.sub(r'<!DOCTYPE[^>]*>', '', html)

    def replace_en_media(m: re.Match) -> str:
        tag = m.group(0)
        hash_match = re.search(r'hash="([0-9a-f]+)"', tag)
        type_match = re.search(r'type="([^"]+)"', tag)
        if not hash_match or not type_match:
            return ''
        h, mime = hash_match.group(1), type_match.group(1)
        att = hash_map.get(h)
        if att and mime in _EMBEDDABLE_IMAGE_MIME:
            b64 = base64.b64encode(att.data).decode()
            width_match = re.search(r'\bwidth="(\d+)"', tag)
            w = int(width_match.group(1)) if width_match else None
            _MAX_WIDTH = 500  # px — fits a standard docx page with margins
            if w is None or w > _MAX_WIDTH:
                w = _MAX_WIDTH
            return f'<img src="data:{mime};base64,{b64}" width="{w}px"/>'
        return ''  # non-image or unknown hash — drop it

    html = re.sub(r'<en-media\b[^>]*/>', replace_en_media, html)
    html = re.sub(r'<en-media\b[^>]*>.*?</en-media>', replace_en_media, html, flags=re.DOTALL)
    # Remove encrypted blocks
    html = re.sub(r'<en-crypt\b[^>]*>.*?</en-crypt>', '', html, flags=re.DOTALL)
    # Convert checkboxes to text markers
    html = re.sub(r'<en-todo\b[^>]*checked="true"[^>]*/>', '[x]\u00a0', html)
    html = re.sub(r'<en-todo\b[^>]*/>', '[\u00a0]\u00a0', html)
    # Strip the ENML root tag (unknown tags confuse html4docx)
    html = re.sub(r'</?en-note\b[^>]*>', '', html)
    return html.strip()




def _compact_doc_spacing(doc: Document) -> None:
    """Remove default paragraph space-after from the document defaults.

    The python-docx template ships with after=200 twips (10 pt) and line=276
    (1.15×). Evernote renders its ENML as plain HTML where <div> blocks have no
    inter-paragraph gap. Override the docDefaults so the DOCX matches.
    """
    styles_elem = doc.part.styles._element
    doc_defaults = styles_elem.find(qn("w:docDefaults"))
    if doc_defaults is None:
        return
    pPr_default = doc_defaults.find(qn("w:pPrDefault"))
    if pPr_default is None:
        return
    pPr = pPr_default.find(qn("w:pPr"))
    if pPr is None:
        return
    spacing = pPr.find(qn("w:spacing"))
    if spacing is not None:
        spacing.set(qn("w:after"), "0")
        spacing.set(qn("w:line"), "240")
        spacing.set(qn("w:lineRule"), "auto")


def _build_doc(
    title: str,
    note: Note,
    attachments: list[Attachment],
    folder: Path,
) -> Document:
    """
    Build a python-docx Document from the note's ENML, preserving formatting.
    Images are embedded inline; PDFs are written as sibling files and referenced by name.
    Returns the Document object (caller saves it).
    """
    doc = Document()
    _compact_doc_spacing(doc)
    h2d = HtmlToDocx()

    hash_map = _attachment_hash_map(attachments)

    parts = []
    if note.source_url:
        url = note.source_url
        parts.append(f'<p>Source: <a href="{url}">{url}</a></p>')
    if note.enml:
        parts.append(_sanitize_enml(note.enml, hash_map))

    html = '\n'.join(parts)
    if html.strip():
        h2d.add_html_to_document(html, doc)

    # Post-process paragraphs
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import nsmap
    _pic_tag = qn('a:graphicData')
    for para in doc.paragraphs:
        if _is_rtl(para.text):
            _set_para_rtl(para)
        # Center paragraphs that contain only an inline image and no text
        if not para.text.strip() and any(r._r.find('.//' + qn('a:graphicData')) is not None for r in para.runs):
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Mirror <w:sz> into <w:szCs> on every run that has an explicit font size.
        # LibreOffice (and Word) use szCs for complex/bidi scripts (Hebrew, Arabic).
        # html4docx only sets sz; without a matching szCs the run inherits the
        # document default (11 pt), so Hebrew text always appears at 11 pt regardless
        # of the sz value.
        for run in para.runs:
            rpr = run._r.find(qn('w:rPr'))
            if rpr is None:
                continue
            sz = rpr.find(qn('w:sz'))
            if sz is None:
                continue
            szcs = rpr.find(qn('w:szCs'))
            if szcs is None:
                szcs = OxmlElement('w:szCs')
                szcs.set(qn('w:val'), sz.get(qn('w:val')))
                sz.addnext(szcs)
            else:
                szcs.set(qn('w:val'), sz.get(qn('w:val')))

    # Non-image attachments (PDFs etc.) → sibling files + hyperlink in doc
    # Images are already embedded inline via <img> tags in the HTML above.
    for i, att in enumerate(attachments, start=1):
        if att.mime not in _EMBEDDABLE_IMAGE_MIME:
            filename = attachment_drive_filename(title, i, att)
            sibling = _unique_path(folder / filename)
            sibling.write_bytes(att.data)
            _set_timestamps(sibling, note.created, note.updated)
            _add_file_hyperlink(doc, f"[Attachment: {sibling.name}]", sibling.name)

    return doc


# ── public API ────────────────────────────────────────────────────────────────

def write_note(
    classified: ClassifiedNote,
    output_dir: Path,
    multi_attachment: str,  # "doc" | "files"
) -> list[Path]:
    """
    Write the note to disk. Returns list of paths created.
    """
    note = classified.note
    folder = note_folder(output_dir, note)
    safe_title = _safe_name(note.title)

    if classified.kind == NoteKind.ATTACHMENT_ONLY_SINGLE:
        att = note.attachments[0]
        ext = _ext_for_mime(att.mime)
        dest = _unique_path(folder / f"{safe_title}{ext}")
        dest.write_bytes(att.data)
        _set_timestamps(dest, note.created, note.updated)
        return [dest]

    elif classified.kind == NoteKind.ATTACHMENT_ONLY_MULTI:
        if multi_attachment == "files":
            paths = []
            for i, att in enumerate(note.attachments, start=1):
                filename = attachment_drive_filename(safe_title, i, att)
                dest = _unique_path(folder / filename)
                dest.write_bytes(att.data)
                _set_timestamps(dest, note.created, note.updated)
                paths.append(dest)
            return paths
        else:  # doc
            doc = _build_doc(note.title, note, note.attachments, folder)
            has_siblings = any(att.mime not in _EMBEDDABLE_IMAGE_MIME for att in note.attachments)
            docx_name = f"{safe_title}_0.docx" if has_siblings else f"{safe_title}.docx"
            dest = _unique_path(folder / docx_name)
            doc.save(str(dest))

            _set_timestamps(dest, note.created, note.updated)
            return [dest]

    elif classified.kind == NoteKind.TEXT_ONLY:
        doc = _build_doc(note.title, note, [], folder)
        dest = _unique_path(folder / f"{safe_title}.docx")
        doc.save(str(dest))
        _set_timestamps(dest, note.created, note.updated)
        return [dest]

    elif classified.kind == NoteKind.TEXT_WITH_ATTACHMENTS:
        doc = _build_doc(note.title, note, note.attachments, folder)
        has_siblings = any(att.mime not in _EMBEDDABLE_IMAGE_MIME for att in note.attachments)
        docx_name = f"{safe_title}_0.docx" if has_siblings else f"{safe_title}.docx"
        dest = _unique_path(folder / docx_name)
        doc.save(str(dest))
        _set_timestamps(dest, note.created, note.updated)
        return [dest]

    raise ValueError(f"Unhandled note kind: {classified.kind}")
